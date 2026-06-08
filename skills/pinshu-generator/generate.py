#!/usr/bin/env python3
"""聘书批量生成 — 核心脚本。SKILL.md 引用此文件。"""

import os, copy, re, zipfile
from docx import Document
from docx.oxml.ns import qn
from lxml import etree
from collections import defaultdict

DESKTOP = os.path.join(os.path.expanduser("~"), "Desktop")

def _dedup(s):
    if len(s) <= 2: return s
    for div in [2, 3]:
        n = len(s) // div
        if n > 0 and all(s[i*n:(i+1)*n] == s[:n] for i in range(div)):
            return s[:n]
    return s

def _cell_text(cell):
    texts = []
    for p in cell.findall(qn('w:p')):
        pt = ''
        for t_el in p.findall('.//' + qn('w:t')):
            if t_el.text: pt += _dedup(t_el.text)
        pt = _dedup(pt.strip())
        if pt and pt not in texts: texts.append(pt)
    return '\n'.join(texts)

def _date_match(text):
    m = re.search(r'(\d+月\s*\d+日)', text)
    return m.group(1).replace(' ', '') if m else None

def _period(text):
    if '上午' in text: return '上午'
    if '下午' in text: return '下午'
    tm = re.search(r'(\d+):\d+\s*-\s*(\d+):\d+', text)
    if tm: return '上午' if int(tm.group(2)) <= 12 else '下午'
    return ''

def _parse_name(text):
    name, prev = text, ''
    while name != prev:
        prev, name = name, re.sub(r'（[^（）]*）', '', name).strip()
    return name or re.sub(r'[（()）]', '', text).strip()

def _parse_expert(text):
    lines = [l.strip() for l in text.split('\n') if l.strip()]
    if not lines: return '', ''
    name = _parse_name(lines[0])
    unit = ''
    for l in lines[1:]:
        if re.match(r'^1[3-9]\d{9}$', l.strip()): continue
        unit += l
    return name, unit

def parse_schedule(filepath):
    doc = Document(filepath)
    rows_el = doc.tables[0]._tbl.findall('.//' + qn('w:tr'))
    experts, carry_date, carry_course = [], '', ''
    prev_expert = None
    for row_el in rows_el[1:]:
        cells = row_el.findall(qn('w:tc'))
        if len(cells) < 4: continue
        dc, vm_date = cells[0], cells[0].find('.//' + qn('w:vMerge'))
        if vm_date is not None:
            if vm_date.get(qn('w:val')) == 'restart':
                raw = _cell_text(dc)
                carry_date = _date_match(raw) or raw.split('\n')[0].strip()
        else:
            dm = _date_match(_cell_text(dc))
            if dm: carry_date = dm
        time_text = _cell_text(cells[1])
        period_val = _period(time_text)
        cc, vm_course = cells[2], cells[2].find('.//' + qn('w:vMerge'))
        course_raw = _cell_text(cc)
        if vm_course is not None:
            course = carry_course if not vm_course.get(qn('w:val')) else course_raw
            if vm_course.get(qn('w:val')) == 'restart': carry_course = course_raw
        else:
            course = course_raw
            if course: carry_course = course
        ec, vm_exp = cells[3], cells[3].find('.//' + qn('w:vMerge'))
        expert_text = _cell_text(ec)
        skip = not course or not time_text or any(
            w in course or w in time_text for w in ['报到', '开班', '案例讲解'])
        if vm_exp is not None:
            if vm_exp.get(qn('w:val')) == 'restart':
                name, unit = _parse_expert(expert_text)
                if name and not skip:
                    prev_expert = (name, unit)
                    experts.append({'name': name, 'unit': unit, 'course': course,
                                   'date': carry_date, 'period': period_val})
            elif prev_expert and not skip:
                experts.append({'name': prev_expert[0], 'unit': prev_expert[1],
                               'course': course, 'date': carry_date, 'period': period_val})
        else:
            name, unit = _parse_expert(expert_text)
            if name and not skip:
                prev_expert = (name, unit)
                experts.append({'name': name, 'unit': unit, 'course': course,
                               'date': carry_date, 'period': period_val})
    return experts

def merge_teachers(experts):
    groups = defaultdict(lambda: {'name': '', 'unit': '', 'sessions': [], 'courses': []})
    for e in experts:
        g = groups[e['name'] + e['unit']]
        g['name'], g['unit'] = e['name'], e['unit']
        g['sessions'].append({'date': e['date'], 'period': e['period']})
        course_lines = [l.strip() for l in re.split(r'[|\n]', e['course']) if l.strip()]
        merged = []
        for cl in course_lines:
            if re.match(r'^\d+[、.）)]', cl): merged.append(cl)
            elif merged: merged[-1] += ' ' + cl
            else: merged.append(cl)
        clean = '\n'.join(merged)
        if clean not in g['courses']: g['courses'].append(clean)
    result = []
    for g in groups.values():
        day_map = defaultdict(set)
        for s in g['sessions']: day_map[s['date']].add(s['period'])
        parts = []
        for d in sorted(day_map.keys(), key=lambda d: [int(x) for x in re.findall(r'\d+', d)]):
            ps = day_map[d]
            parts.append(d + ('全天' if '上午' in ps and '下午' in ps
                         else '上午' if '上午' in ps
                         else '下午' if '下午' in ps else ''))
        result.append({'name': g['name'], 'unit': g['unit'],
                       'time': '、'.join(parts), 'course': '\n'.join(g['courses'])})
    return result

def fill_doc(doc, teacher, common):
    repl = {
        '（会议名字）': common['meeting_name'], '（对方单位）': teacher['unit'],
        '（起止时间）': common['start_date'], '（结束时间）': common['end_date'],
        '（地点）': common['location'], '（老师名字）': teacher['name'],
        '（上课时间）': teacher['time'], '（课程内容）': '__COURSE__',
        '（授课地址）': common['address'], '（发函时间）': common['letter_date'],
    }
    for p in doc.paragraphs:
        for r in p.runs:
            for k, v in repl.items():
                if k in r.text: r.text = r.text.replace(k, v)
    lines = teacher['course'].split('\n')
    ci = next((i for i, p in enumerate(doc.paragraphs) if '__COURSE__' in p.text), None)
    if ci is not None and lines:
        cp = doc.paragraphs[ci]
        for r in cp.runs:
            if '__COURSE__' in r.text: r.text = r.text.replace('__COURSE__', lines[0])
        pel = cp._element
        for line in lines[1:]:
            if not line.strip(): continue
            nel = copy.deepcopy(pel)
            for tel in nel.findall('.//' + qn('w:t')): tel.text = ''
            ft = nel.find('.//' + qn('w:t'))
            if ft is not None:
                ft.text = line
                ft.attrib[qn('xml:space')] = 'preserve'
            pel.addnext(nel)
            pel = nel

def generate(template_path, schedule_path, common, output_path=None):
    """一站生成：解析日程 + 填充模板 + 合并Word"""
    teachers = merge_teachers(parse_schedule(schedule_path))
    base = Document(template_path)
    fill_doc(base, teachers[0], common)
    body = base.element.body
    for t in teachers[1:]:
        pp = etree.SubElement(body, qn('w:p'))
        pr = etree.SubElement(pp, qn('w:r'))
        etree.SubElement(pr, qn('w:br')).set(qn('w:type'), 'page')
        tmp = Document(template_path)
        fill_doc(tmp, t, common)
        for child in list(tmp.element.body):
            body.append(child)
    if not output_path:
        output_path = os.path.join(DESKTOP, '聘书_全部.docx')
    base.save(output_path)
    return output_path, len(teachers)
